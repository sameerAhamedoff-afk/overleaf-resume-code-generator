import argparse
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_cover_letter(output_path, name, email, phone, linkedin, company, role, content_paragraphs):
    doc = Document()
    
    # Page setup - Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Contact Header (Centered, bold name, clean details)
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_header.add_run(name + "\n")
    run_name.font.size = Pt(18)
    run_name.bold = True
    
    details_str = f"Phone: {phone} | Email: {email} | LinkedIn: {linkedin}\n"
    run_details = p_header.add_run(details_str)
    run_details.font.size = Pt(10)
    
    # Add a thin divider line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line_paragraph_format = p_line.paragraph_format
    p_line_paragraph_format.space_after = Pt(24)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '555555')
    pBdr.append(bottom)
    p_line._p.get_or_add_pPr().append(pBdr)

    # Date
    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_after = Pt(12)
    p_date.add_run(datetime.now().strftime("%B %d, %Y"))

    # Recipient Info
    p_recipient = doc.add_paragraph()
    p_recipient.paragraph_format.space_after = Pt(18)
    p_recipient.add_run(f"Hiring Team\n{company}\n")

    # Salutation
    p_salutation = doc.add_paragraph()
    p_salutation.paragraph_format.space_after = Pt(12)
    p_salutation.add_run(f"Dear Hiring Team at {company},")

    # Content Paragraphs
    for para in content_paragraphs:
        if para.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.15
            p.add_run(para.strip())

    # Sign-off
    p_signoff = doc.add_paragraph()
    p_signoff.paragraph_format.space_before = Pt(18)
    p_signoff.paragraph_format.space_after = Pt(24)
    p_signoff.add_run("Sincerely,\n\n")
    
    p_name = doc.add_paragraph()
    run_sign_name = p_name.add_run(name)
    run_sign_name.bold = True
    
    # Save document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Cover letter successfully created at: {output_path}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Generate Word Cover Letter")
    parser.add_argument("--output", required=True, help="Output docx file path")
    parser.add_argument("--name", required=True, help="Applicant Name")
    parser.add_argument("--email", required=True, help="Applicant Email")
    parser.add_argument("--phone", required=True, help="Applicant Phone")
    parser.add_argument("--linkedin", required=True, help="Applicant LinkedIn")
    parser.add_argument("--company", required=True, help="Target Company")
    parser.add_argument("--role", required=True, help="Target Role")
    parser.add_argument("--content_file", required=True, help="Path to text file containing cover letter paragraphs")
    
    args = parser.parse_args()
    
    with open(args.content_file, 'r', encoding='utf-8-sig') as f:
        paragraphs = f.read().split('\n\n')
        
    generate_cover_letter(
        output_path=args.output,
        name=args.name,
        email=args.email,
        phone=args.phone,
        linkedin=args.linkedin,
        company=args.company,
        role=args.role,
        content_paragraphs=paragraphs
    )
