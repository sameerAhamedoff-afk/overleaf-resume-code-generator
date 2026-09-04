import argparse
import os
import json
import subprocess
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def add_bottom_border(paragraph, border_type="single", sz="18", space="1", color="0070c0"):
    """Adds a bottom border to a paragraph using OpenXML."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(r'<w:pBdr %s><w:bottom w:val="%s" w:sz="%s" w:space="%s" w:color="%s"/></w:pBdr>' % (nsdecls('w'), border_type, sz, space, color))
    pPr.append(pBdr)

def add_hyperlink(paragraph, url, text, color="0000ff", underline=True):
    """Adds a real hyperlink element to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = parse_xml(r'<w:hyperlink %s xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:id="%s"/>' % (nsdecls('w'), r_id))
    new_run = parse_xml(r'<w:r %s/>' % nsdecls('w'))
    text_node = parse_xml(r'<w:t %s>%s</w:t>' % (nsdecls('w'), text))
    new_run.append(text_node)

    rPr = parse_xml(r'<w:rPr %s/>' % nsdecls('w'))
    rFonts = parse_xml(r'<w:rFonts %s w:ascii="Calibri" w:hAnsi="Calibri" w:cs="Calibri"/>' % nsdecls('w'))
    rPr.append(rFonts)

    b = parse_xml(r'<w:b %s w:val="1"/>' % nsdecls('w'))
    rPr.append(b)

    sz = parse_xml(r'<w:sz %s w:val="20"/>' % nsdecls('w'))
    rPr.append(sz)

    c = parse_xml(r'<w:color %s w:val="%s"/>' % (nsdecls('w'), color))
    rPr.append(c)

    if underline:
        u = parse_xml(r'<w:u %s w:val="single"/>' % nsdecls('w'))
        rPr.append(u)

    new_run.append(rPr)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def clean_project_name(name, company_name):
    """Strips the company name prefix from the project title."""
    for sep in [" — ", " – ", " - ", " | "]:
        if sep in name:
            parts = name.split(sep, 1)
            company_words = set(w.lower() for w in company_name.split() if len(w) > 2)
            left_words = set(w.lower() for w in parts[0].split())
            if company_words.intersection(left_words):
                return parts[1].strip()
    
    clean = name
    if company_name.lower() in clean.lower():
        idx = clean.lower().find(company_name.lower())
        clean = clean[idx + len(company_name):].strip()
        clean = clean.lstrip(" -–—|:").strip()
    return clean

def generate_resume(data_path, output_path, include_certifications=None):
    with open(data_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Determine whether to include certifications
    if include_certifications is None:
        if "include_certifications" in data:
            include_certifications = bool(data["include_certifications"])
        elif "enable_certifications" in data:
            include_certifications = bool(data["enable_certifications"])
        elif "certifications_enabled" in data:
            include_certifications = bool(data["certifications_enabled"])
        else:
            include_certifications = True

    doc = Document()

    # Section Setup - Page margins & size (A4)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.375)
    section.bottom_margin = Inches(0.20)
    section.left_margin = Inches(0.394)
    section.right_margin = Inches(0.487)

    # Styles Setup (Calibri, 10pt standard)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    def sanitize_text(text):
        if not text:
            return ""
        for symbol in ["☎:", "☎", "🖂:", "🖂", "📍:", "📍", "✉:", "✉", "📞:", "📞"]:
            text = text.replace(symbol, "")
        return text.strip()

    # 1. Document Header (Name and Target Role / Headline)
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(1)
    p_name.paragraph_format.line_spacing = 1.0

    run_name = p_name.add_run(data["name"].upper())
    run_name.font.name = 'Calibri'
    run_name.font.size = Pt(18)
    run_name.bold = True

    role_title = data.get("headline") or data.get("role") or ""
    if role_title:
        p_role_title = doc.add_paragraph()
        p_role_title.paragraph_format.space_before = Pt(0)
        p_role_title.paragraph_format.space_after = Pt(2)
        p_role_title.paragraph_format.line_spacing = 1.0

        run_role = p_role_title.add_run(role_title)
        run_role.font.name = 'Calibri'
        run_role.font.size = Pt(12)
        run_role.bold = True

    # 2. Contact Details (Clean, ATS-friendly without unicode emojis)
    p_contact = doc.add_paragraph()
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(4)
    p_contact.paragraph_format.line_spacing = 1.0

    contact_parts = []
    phone = sanitize_text(data.get("phone", ""))
    email = sanitize_text(data.get("email", ""))
    location = sanitize_text(data.get("location", ""))

    if phone:
        contact_parts.append(phone)
    if email:
        contact_parts.append(email)
    if location:
        contact_parts.append(location)

    contact_str = " | ".join(contact_parts)
    if contact_str:
        run_contact = p_contact.add_run(contact_str)
        run_contact.font.name = 'Calibri'
        run_contact.font.size = Pt(10)
        run_contact.bold = True
        run_contact.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

    if data.get("linkedin"):
        if contact_parts:
            run_sep = p_contact.add_run(" | ")
            run_sep.font.name = 'Calibri'
            run_sep.font.size = Pt(10)
            run_sep.bold = True
            run_sep.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
        add_hyperlink(p_contact, data["linkedin"], "LinkedIn", color="0000ff", underline=True)

    if data.get("github"):
        run_sep_gh = p_contact.add_run(" | ")
        run_sep_gh.font.name = 'Calibri'
        run_sep_gh.font.size = Pt(10)
        run_sep_gh.bold = True
        run_sep_gh.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
        add_hyperlink(p_contact, data["github"], "GitHub", color="0000ff", underline=True)

    add_bottom_border(p_contact, "single", "6", "1", "d0d0d0")

    def add_section_heading(text):
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(8)
        p_head.paragraph_format.space_after = Pt(3)
        p_head.paragraph_format.line_spacing = 1.1
        run_head = p_head.add_run(text.upper())
        run_head.font.name = 'Calibri'
        run_head.font.size = Pt(11)
        run_head.bold = True
        run_head.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
        add_bottom_border(p_head, "single", "18", "1", "0070c0")
        return p_head

    # 3. Professional Summary (Explicit section heading for ATS detection)
    if data.get("summary"):
        add_section_heading("PROFESSIONAL SUMMARY")
        p_summary = doc.add_paragraph()
        p_summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_summary.paragraph_format.space_before = Pt(0)
        p_summary.paragraph_format.space_after = Pt(4)
        p_summary.paragraph_format.line_spacing = 1.1
        run_summary = p_summary.add_run(data["summary"])
        run_summary.font.name = 'Calibri'
        run_summary.font.size = Pt(10)

    # 4. Technical Skills (ATS-optimized: categorized, comma-separated, NO pipe walls!)
    add_section_heading("TECHNICAL SKILLS")
    tech_skills_dict = {}
    if "technical_skills" in data and isinstance(data["technical_skills"], dict):
        tech_skills_dict = data["technical_skills"]
    elif "tools" in data and isinstance(data["tools"], dict):
        tech_skills_dict = data["tools"]
    elif "skills" in data and isinstance(data["skills"], dict):
        tech_skills_dict = data["skills"]
    elif "skills" in data and isinstance(data["skills"], list):
        for item in data["skills"]:
            if ":" in item:
                cat, val = item.split(":", 1)
                tech_skills_dict[cat.strip()] = val.strip()
            else:
                cleaned_item = item.replace(" | ", ", ").replace("|", ",")
                tech_skills_dict.setdefault("Core Skills", []).append(cleaned_item)

    for key, val in tech_skills_dict.items():
        p_tool = doc.add_paragraph()
        p_tool.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_tool.paragraph_format.space_before = Pt(0)
        p_tool.paragraph_format.space_after = Pt(2)
        p_tool.paragraph_format.line_spacing = 1.1

        run_cat = p_tool.add_run(f"{key}: ")
        run_cat.font.name = 'Calibri'
        run_cat.font.size = Pt(10)
        run_cat.bold = True

        if isinstance(val, list):
            val_str = ", ".join(str(v).strip() for v in val)
        else:
            val_str = str(val)

        # Eliminate any pipeline characters inside values to avoid confusing linear ATS parsers
        val_str = val_str.replace(" | ", ", ").replace(" |", ",").replace("| ", ", ")

        run_val = p_tool.add_run(val_str)
        run_val.font.name = 'Calibri'
        run_val.font.size = Pt(10)

    # 5. Professional Experience
    add_section_heading("PROFESSIONAL EXPERIENCE")
    
    # Associate projects to roles dynamically using multi-stage matching if projects exist
    role_projects = {role['company']: [] for role in data.get("roles", [])}
    unassigned_projects = []
    
    for proj in data.get("projects", []):
        assigned = False
        proj_name_lower = proj.get('name', '').lower()
        proj_company_lower = proj.get("company", "").lower()
        bullets_text_lower = " ".join(proj.get("bullets", [])).lower()
        
        if "tios" in proj_name_lower:
            proj_name_lower = proj_name_lower.replace("tios", "trios")
        if "tios" in proj_company_lower:
            proj_company_lower = proj_company_lower.replace("tios", "trios")
        if "tios" in bullets_text_lower:
            bullets_text_lower = bullets_text_lower.replace("tios", "trios")

        for role in data.get("roles", []):
            role_company_lower = role.get('company', '').lower()
            if "tios" in role_company_lower:
                role_company_lower = role_company_lower.replace("tios", "trios")
                
            if proj_company_lower and (proj_company_lower in role_company_lower or role_company_lower in proj_company_lower):
                role_projects[role['company']].append(proj)
                assigned = True
                break
                
            company_words = [w.lower() for w in role['company'].replace("tios", "trios").split() if len(w) > 2 and w.lower() not in ["pvt", "ltd"]]
            if company_words and any(w in proj_name_lower for w in company_words):
                role_projects[role['company']].append(proj)
                assigned = True
                break
                
            if company_words and any(w in bullets_text_lower for w in company_words):
                role_projects[role['company']].append(proj)
                assigned = True
                break
                
        if not assigned:
            unassigned_projects.append(proj)
            
    if unassigned_projects and data.get("roles"):
        role_projects[data["roles"][0]['company']].extend(unassigned_projects)

    for role in data.get("roles", []):
        # Role Header (Clean 2-line layout for flawless ATS linear parsing)
        p_role = doc.add_paragraph()
        p_role.paragraph_format.space_before = Pt(6)
        p_role.paragraph_format.space_after = Pt(1)
        p_role.paragraph_format.line_spacing = 1.15

        title_comp = f"{role.get('title', '')} | {role.get('company', '')}"
        run_role_title = p_role.add_run(title_comp)
        run_role_title.font.name = 'Calibri'
        run_role_title.font.size = Pt(10)
        run_role_title.bold = True

        p_loc = doc.add_paragraph()
        p_loc.paragraph_format.space_before = Pt(0)
        p_loc.paragraph_format.space_after = Pt(2)
        p_loc.paragraph_format.line_spacing = 1.1

        loc_timeline = f"{role.get('location', '')} | {role.get('timeline', '')}"
        run_loc_time = p_loc.add_run(loc_timeline)
        run_loc_time.font.name = 'Calibri'
        run_loc_time.font.size = Pt(9.5)
        run_loc_time.font.italic = True
        run_loc_time.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Direct role bullets
        for bullet in role.get("bullets", []):
            p_bullet = doc.add_paragraph()
            p_bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_bullet.paragraph_format.left_indent = Inches(0.25)
            p_bullet.paragraph_format.first_line_indent = Inches(-0.25)
            p_bullet.paragraph_format.space_before = Pt(0)
            p_bullet.paragraph_format.space_after = Pt(1)
            p_bullet.paragraph_format.line_spacing = 1.15

            pPr = p_bullet._p.get_or_add_pPr()
            numPr = parse_xml(r'<w:numPr %s><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>' % nsdecls('w'))
            pPr.append(numPr)

            run_bullet = p_bullet.add_run(bullet)
            run_bullet.font.name = 'Calibri'
            run_bullet.font.size = Pt(10)

        # Projects / domains under this role (if any)
        projects_for_role = role_projects.get(role.get('company', ''), [])
        for proj in projects_for_role:
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_title.paragraph_format.space_before = Pt(3)
            p_title.paragraph_format.space_after = Pt(1)
            p_title.paragraph_format.line_spacing = 1.1

            clean_name = clean_project_name(proj.get("name", ""), role.get('company', ''))
            run_title = p_title.add_run(clean_name)
            run_title.font.name = 'Calibri'
            run_title.font.size = Pt(10)
            run_title.bold = True

            for bullet in proj.get("bullets", []):
                p_bullet = doc.add_paragraph()
                p_bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_bullet.paragraph_format.left_indent = Inches(0.25)
                p_bullet.paragraph_format.first_line_indent = Inches(-0.25)
                p_bullet.paragraph_format.space_before = Pt(0)
                p_bullet.paragraph_format.space_after = Pt(1)
                p_bullet.paragraph_format.line_spacing = 1.15

                pPr = p_bullet._p.get_or_add_pPr()
                numPr = parse_xml(r'<w:numPr %s><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>' % nsdecls('w'))
                pPr.append(numPr)

                run_bullet = p_bullet.add_run(bullet)
                run_bullet.font.name = 'Calibri'
                run_bullet.font.size = Pt(10)

    # 6. Education
    add_section_heading("EDUCATION")
    for edu in data.get("education", []):
        p_deg = doc.add_paragraph()
        p_deg.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_deg.paragraph_format.space_before = Pt(3)
        p_deg.paragraph_format.space_after = Pt(1)
        p_deg.paragraph_format.line_spacing = 1.15

        run_deg = p_deg.add_run(edu.get("degree", ""))
        run_deg.font.name = 'Calibri'
        run_deg.font.size = Pt(10)
        run_deg.bold = True

        inst_timeline_parts = []
        if edu.get("institution"):
            inst_timeline_parts.append(edu["institution"])
        if edu.get("timeline"):
            inst_timeline_parts.append(edu["timeline"])
        elif edu.get("year"):
            inst_timeline_parts.append(edu["year"])

        if inst_timeline_parts:
            p_inst = doc.add_paragraph()
            p_inst.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_inst.paragraph_format.space_before = Pt(0)
            p_inst.paragraph_format.space_after = Pt(2)
            p_inst.paragraph_format.line_spacing = 1.1

            run_inst = p_inst.add_run(" | ".join(inst_timeline_parts))
            run_inst.font.name = 'Calibri'
            run_inst.font.size = Pt(9.5)
            run_inst.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 7. Certifications
    if include_certifications and data.get("certifications"):
        add_section_heading("CERTIFICATIONS")
        for cert in data["certifications"]:
            p_cert = doc.add_paragraph()
            p_cert.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_cert.paragraph_format.left_indent = Inches(0.25)
            p_cert.paragraph_format.first_line_indent = Inches(-0.25)
            p_cert.paragraph_format.space_before = Pt(0)
            p_cert.paragraph_format.space_after = Pt(1)
            p_cert.paragraph_format.line_spacing = 1.15

            pPr = p_cert._p.get_or_add_pPr()
            numPr = parse_xml(r'<w:numPr %s><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>' % nsdecls('w'))
            pPr.append(numPr)

            if isinstance(cert, dict):
                cert_name = cert.get("name", "")
                cert_org = cert.get("org", "")
            else:
                cert_name = str(cert)
                cert_org = ""

            run_cname = p_cert.add_run(cert_name)
            run_cname.font.name = 'Calibri'
            run_cname.font.size = Pt(10)
            run_cname.bold = True

            if cert_org:
                run_corg = p_cert.add_run(f" | {cert_org}")
                run_corg.font.name = 'Calibri'
                run_corg.font.size = Pt(10)

    # Save DOCX
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)
    print(f"Word resume successfully generated at: {output_path}")

def convert_docx_to_pdf(docx_path, pdf_path):
    """Converts a DOCX file to PDF using Microsoft Word COM interface via PowerShell."""
    docx_path = os.path.abspath(docx_path)
    pdf_path = os.path.abspath(pdf_path)

    # Temporary ps1 file path
    temp_dir = os.path.dirname(pdf_path)
    ps1_path = os.path.join(temp_dir, "temp_convert_pdf.ps1")

    ps1_content = f"""
    $clean_path = "{docx_path}"
    $pdf_path = "{pdf_path}"

    try {{
        $word = New-Object -ComObject Word.Application
        $word.DisplayAlerts = 0
        $word.Visible = $false
        $doc = $word.Documents.Open($clean_path, $false, $true, $false)
        $doc.SaveAs($pdf_path, 17)
        $doc.Close()
        $word.Quit()
        [System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null
        [System.GC]::Collect()
        [System.GC]::WaitForPendingFinalizers()
    }} catch {{
        Write-Error $_.Exception.Message
    }}
    """

    with open(ps1_path, 'w', encoding='utf-8') as f:
        f.write(ps1_content)

    print("Running PDF compilation via background PowerShell script...")
    result = subprocess.run(["powershell", "-ExecutionPolicy", "Bypass", "-File", ps1_path], capture_output=True, text=True)

    if os.path.exists(ps1_path):
        try:
            os.remove(ps1_path)
        except Exception:
            pass

    if os.path.exists(pdf_path):
        print(f"PDF resume successfully generated at: {pdf_path}")
    else:
        print(f"Failed to generate PDF! PowerShell Error: {result.stderr}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Generate Word & PDF Resume (Template-4 Layout)")
    parser.add_argument("--json", required=True, help="Path to JSON file containing resume data")
    parser.add_argument("--output", required=True, help="Output .docx file path")
    parser.add_argument("--pdf", required=False, help="Optional output .pdf file path")

    cert_group = parser.add_mutually_exclusive_group()
    cert_group.add_argument(
        "--certifications", "--include-certifications",
        dest="include_certifications",
        action="store_true",
        default=None,
        help="Explicitly enable/include certifications in the generated resume"
    )
    cert_group.add_argument(
        "--no-certifications", "--exclude-certifications", "--avoid-certifications",
        dest="include_certifications",
        action="store_false",
        help="Avoid mentioning/exclude certifications in the generated resume"
    )

    args = parser.parse_args()
    generate_resume(args.json, args.output, include_certifications=args.include_certifications)

    if args.pdf:
        with open(args.json, 'r', encoding='utf-8') as f:
            data = json.load(f)
        role = data.get("role", "resume")
        role_clean = role.replace(" ", "_").replace("/", "_").replace("-", "_")
        role_clean = "".join(c for c in role_clean if c.isalnum() or c == "_")
        while "__" in role_clean:
            role_clean = role_clean.replace("__", "_")
        
        pdf_dir = os.path.dirname(args.pdf) or "."
        if args.pdf.lower().endswith(".pdf") and os.path.basename(args.pdf).lower() != "resume.pdf":
            pdf_path = os.path.abspath(args.pdf)
        else:
            pdf_path = os.path.join(pdf_dir, f"Sameer_{role_clean}.pdf")
        
        # Clean up any generic resume.pdf that might be left over
        old_pdf = os.path.join(pdf_dir, "resume.pdf")
        if os.path.exists(old_pdf) and os.path.abspath(old_pdf) != os.path.abspath(pdf_path):
            try:
                os.remove(old_pdf)
            except Exception:
                pass
                
        convert_docx_to_pdf(args.output, pdf_path)
