import os
import sys
import re
import json
import argparse
from urllib.parse import urlparse

# Import docx if available for Word resume support
try:
    import docx
except ImportError:
    docx = None

def strip_html_tags(html_content):
    """Regex-based HTML tag stripper to convert HTML page structures to plain text."""
    # Remove script and style elements
    html_content = re.sub(r'<script[^>]*>([\s\S]*?)</script>', '', html_content)
    html_content = re.sub(r'<style[^>]*>([\s\S]*?)</style>', '', html_content)
    # Remove HTML comments
    html_content = re.sub(r'<!--([\s\S]*?)-->', '', html_content)
    # Replace tag closures/brs with newlines
    html_content = re.sub(r'<br\s*/?>', '\n', html_content)
    html_content = re.sub(r'</?(?:p|div|h\d|li|ul|ol|tr|td|table|section)[^>]*>', '\n', html_content)
    # Strip all remaining tags
    text = re.sub(r'<[^>]+>', ' ', html_content)
    # Decode basic HTML entities
    text = text.replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
    # Normalize whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n+', '\n\n', text)
    return text.strip()

def extract_text_from_tex(filepath):
    """Extracts clean text content from a LaTeX (.tex) resume file."""
    if not os.path.exists(filepath):
        print(f"Error: Resume file not found: {filepath}", file=sys.stderr)
        return ""
    
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
        
    # Remove LaTeX comments (handling escaped \%)
    cleaned_lines = []
    for line in content.splitlines():
        cleaned_line = ""
        escaped = False
        for char in line:
            if char == '\\':
                escaped = not escaped
                cleaned_line += char
            elif char == '%':
                if not escaped:
                    break  # comment starts
                else:
                    cleaned_line += char
                    escaped = False
            else:
                escaped = False
                cleaned_line += char
        cleaned_lines.append(cleaned_line)
    text = "\n".join(cleaned_lines)
    
    # Extract only document body if \begin{document} is present
    doc_match = re.search(r'\\begin\{document\}', text)
    if doc_match:
        text = text[doc_match.end():]
    text = re.sub(r'\\end\{document\}', '', text)
    
    # Strip common LaTeX inline formatting/spacing commands but keep their content
    text = re.sub(r'\\(?:item|noindent|hfill|small|Large|large|centering|medskip|smallskip|bigskip|newpage)\b', ' ', text)
    
    # Recursively strip bracket commands like \textbf{...} while keeping inside text
    pattern = r'\\(?:\w+)\{([^{}]*)\}'
    for _ in range(5):  # handle up to 5 nested brace levels
        text = re.sub(pattern, r' \1 ', text)
        
    # Remove remaining unmatched braces
    text = text.replace('{', ' ').replace('}', ' ')
    
    # Remove remaining command symbols and command words
    text = re.sub(r'\\[a-zA-Z]+', ' ', text)
    
    # Unescape escaped special LaTeX characters
    text = text.replace('\\&', '&').replace('\\%', '%').replace('\\_', '_').replace('\\$', '$').replace('\\#', '#')
    return text

def extract_text_from_docx(filepath):
    """Extracts text from a Word (.docx) resume file."""
    if not docx:
        print("Warning: python-docx is not installed. Cannot parse Word resumes.", file=sys.stderr)
        return ""
    if not os.path.exists(filepath):
        print(f"Error: Resume file not found: {filepath}", file=sys.stderr)
        return ""
    
    try:
        doc = docx.Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error reading .docx: {e}", file=sys.stderr)
        return ""

def extract_text_from_file(filepath):
    """Utility to extract raw text based on file format extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.tex':
        return extract_text_from_tex(filepath)
    elif ext == '.docx':
        return extract_text_from_docx(filepath)
    elif ext in ['.md', '.txt']:
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
    else:
        # Fallback reading
        if os.path.exists(filepath):
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception:
                pass
    return ""

def load_skills_from_profile(profile_path):
    """Parses profile.md and dynamically extracts the skills list."""
    skills = set()
    if not os.path.exists(profile_path):
        return skills
    
    try:
        with open(profile_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        lines = content.splitlines()
        in_skills = False
        for line in lines:
            if line.startswith('## Skills Inventory'):
                in_skills = True
                continue
            if in_skills and line.startswith('##'):
                break
            if in_skills and line.strip().startswith('-'):
                # Extract skill items
                match = re.match(r'-\s*\*\*[^*]+\*\*:\s*(.*)', line)
                if match:
                    skills_list = match.group(1)
                else:
                    skills_list = line.replace('-', '').strip()
                
                # Split by comma or pipe
                parts = re.split(r',|\|', skills_list)
                for part in parts:
                    part = part.strip()
                    if not part:
                        continue
                    # Parse parenthetical expressions (e.g. "Large Language Models (LLMs)")
                    p_match = re.match(r'([^(]+)\s*\(([^)]+)\)', part)
                    if p_match:
                        skills.add(p_match.group(1).strip())
                        skills.add(p_match.group(2).strip())
                    else:
                        skills.add(part)
    except Exception as e:
        print(f"Warning parsing profile.md for skills: {e}", file=sys.stderr)
        
    return skills

def get_builtin_keywords():
    """Returns a rich set of fallback/supplementary software engineering and AI terms."""
    return {
        # Languages
        "Python", "JavaScript", "TypeScript", "C++", "C#", "Java", "Go", "Golang", "Rust", "Ruby", "PHP", "SQL", "HTML5", "CSS3", "HTML", "CSS", "R",
        # Frameworks & Libraries
        "React", "React.js", "Next.js", "Angular", "Vue.js", "Vue", "Node.js", "Node", "Express.js", "Express", "FastAPI", "Flask", "Django", "Spring Boot",
        "Tailwind CSS", "Tailwind", "Bootstrap", "jQuery", "Redux", "GraphQL", "PyTorch", "TensorFlow", "scikit-learn", "Keras", "OpenCV", "Hugging Face",
        "LangChain", "LangGraph", "CrewAI", "AutoGen", "LlamaIndex", "Pandas", "NumPy", "SciPy", "Matplotlib", "Seaborn",
        # Cloud & DevOps
        "AWS", "Amazon Web Services", "Azure", "Google Cloud", "GCP", "Docker", "Kubernetes", "K8s", "CI/CD", "GitHub Actions", "GitLab CI", "Jenkins",
        "Terraform", "Ansible", "Nginx", "PM2", "Linux", "Unix", "Bash", "Prometheus", "Grafana", "Datadog", "CloudFormation", "Lambda", "EC2", "S3", "SES",
        # Databases & Cache
        "PostgreSQL", "Postgres", "MySQL", "SQLite", "MongoDB", "NoSQL", "Redis", "Elasticsearch", "Cassandra", "DynamoDB", "Firebase",
        # Vector Databases
        "Pinecone", "Chroma", "ChromaDB", "Weaviate", "Qdrant", "Milvus", "pgvector",
        # AI & Agentic Systems
        "Large Language Models", "LLM", "LLMs", "RAG", "Retrieval-Augmented Generation", "Computer Vision", "Natural Language Processing", "NLP",
        "Sentence Transformers", "Prompt Engineering", "Fine-Tuning", "Structured Outputs", "Tool Calling", "Autonomous Agents", "Agentic Systems",
        "Agentic AI", "AI Agents", "Semantic Search", "Hybrid Search", "BM25", "Re-ranking", "Vector Search", "Evals", "LangSmith", "Ragas", "MLflow",
        # Concepts & Methodologies
        "Microservices", "RESTful APIs", "REST API", "REST APIs", "WebSocket", "WebSockets", "System Design", "OOP", "Object-Oriented Programming",
        "Agile", "Scrum", "TDD", "Test-Driven Development", "Git", "GitHub", "Postman", "Jira", "Software Development Life Cycle", "SDLC", "Data Structures",
        "Algorithms", "JWT", "Authentication", "Responsible AI", "Security", "Governance"
    }

def keyword_exists(text, keyword):
    """Regex-based verification that a keyword appears with word boundaries."""
    escaped = re.escape(keyword)
    pattern = escaped
    # Prefix boundary
    if keyword[0].isalnum():
        pattern = r'(?:^|[^a-zA-Z0-9_])' + pattern
    # Suffix boundary
    if keyword[-1].isalnum():
        pattern = pattern + r'(?:$|[^a-zA-Z0-9_])'
        
    return bool(re.search(pattern, text, re.IGNORECASE))

def calculate_match(resume_text, jd_text, profile_skills):
    """Calculates match metrics between resume and job description."""
    # Combine profile skills and built-in terms to create the vocabulary
    vocab = profile_skills.union(get_builtin_keywords())
    
    # 1. Identify which vocabulary keywords are requested in the Job Description
    required_keywords = []
    # Check longer keywords first to avoid sub-match issues (e.g. "React.js" matching "React")
    for kw in sorted(vocab, key=len, reverse=True):
        if keyword_exists(jd_text, kw):
            required_keywords.append(kw)
            
    # Normalize required keyword list (filter case-variants and keep unique values)
    unique_required = []
    for kw in required_keywords:
        if kw.lower() not in [x.lower() for x in unique_required]:
            unique_required.append(kw)
            
    if not unique_required:
        return {
            "score": 0,
            "matched": [],
            "missing": [],
            "required": []
        }
        
    # 2. Verify presence in the resume
    matched_keywords = []
    missing_keywords = []
    
    for kw in unique_required:
        if keyword_exists(resume_text, kw):
            matched_keywords.append(kw)
        else:
            missing_keywords.append(kw)
            
    score = int((len(matched_keywords) / len(unique_required)) * 100)
    
    return {
        "score": score,
        "matched": matched_keywords,
        "missing": missing_keywords,
        "required": unique_required
    }

def main():
    parser = argparse.ArgumentParser(description="Evaluate match rating between a Resume and a Job Description.")
    parser.add_argument("--resume", required=True, help="Path to resume (.tex, .docx, .md, .txt)")
    parser.add_argument("--jd", required=True, help="Path to JD file, raw JD text, or URL")
    parser.add_argument("--profile", default="profile.md", help="Path to profile.md containing skills inventory")
    parser.add_argument("--json-output", help="Path to write JSON report")
    parser.add_argument("--markdown-output", help="Path to write Markdown report")
    
    args = parser.parse_args()
    
    # Extract Resume Text
    resume_text = extract_text_from_file(args.resume)
    if not resume_text:
        print(f"Error: Could not extract text from resume: {args.resume}", file=sys.stderr)
        sys.exit(1)
        
    # Extract JD Text
    jd_text = ""
    if os.path.exists(args.jd):
        with open(args.jd, 'r', encoding='utf-8') as f:
            jd_text = f.read()
    else:
        parsed_url = urlparse(args.jd)
        if parsed_url.scheme in ['http', 'https']:
            try:
                import requests
                r = requests.get(args.jd, timeout=10)
                if r.status_code == 200:
                    jd_text = strip_html_tags(r.text)
                else:
                    print(f"Error: Failed to fetch JD URL (Status {r.status_code})", file=sys.stderr)
                    sys.exit(1)
            except Exception as e:
                print(f"Error fetching JD URL: {e}", file=sys.stderr)
                sys.exit(1)
        else:
            # Direct text input
            jd_text = args.jd
            
    if not jd_text.strip():
        print("Error: Empty Job Description provided.", file=sys.stderr)
        sys.exit(1)
        
    # Load profile skills
    profile_skills = load_skills_from_profile(args.profile)
    
    # Perform match calculation
    result = calculate_match(resume_text, jd_text, profile_skills)
    
    # Print clean report to console
    print("\n" + "="*60)
    print("                 RESUME MATCH RATING REPORT")
    print("="*60)
    print(f"Resume File:   {os.path.basename(args.resume)}")
    print(f"Match Score:   {result['score']}%")
    print(f"Keywords:      {len(result['matched'])} / {len(result['required'])} matched")
    print("-"*60)
    
    if result['matched']:
        print("MATCHED KEYWORDS:")
        matched_sorted = sorted(result['matched'])
        for i in range(0, len(matched_sorted), 4):
            print("  • " + ", ".join(matched_sorted[i:i+4]))
    else:
        print("MATCHED KEYWORDS: None")
        
    print("-"*60)
    if result['missing']:
        print("MISSING KEYWORDS (Recommended to add):")
        missing_sorted = sorted(result['missing'])
        for i in range(0, len(missing_sorted), 4):
            print("  • " + ", ".join(missing_sorted[i:i+4]))
    else:
        print("MISSING KEYWORDS: None! Perfect match!")
    print("="*60 + "\n")
    
    # Save reports
    if args.json_output:
        os.makedirs(os.path.dirname(os.path.abspath(args.json_output)), exist_ok=True)
        with open(args.json_output, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2)
        print(f"Saved JSON report to: {args.json_output}")
        
    if args.markdown_output:
        os.makedirs(os.path.dirname(os.path.abspath(args.markdown_output)), exist_ok=True)
        
        md_content = f"""# Resume Match Rate Evaluation Report

- **Resume File:** `{os.path.basename(args.resume)}`
- **ATS Match Score:** `{result['score']}%`
- **Keywords Match Count:** `{len(result['matched'])} / {len(result['required'])}`

## Summary Analysis
- A match rate of **{result['score']}%** indicates a {'highly optimized' if result['score'] >= 85 else 'moderately tailored' if result['score'] >= 70 else 'basic/untailored'} alignment with the job requirements.

## Matched Keywords ({len(result['matched'])})
| Keyword / Skill | Status |
| :--- | :--- |
"""
        for kw in sorted(result['matched']):
            md_content += f"| {kw} | Matched ✅ |\n"
            
        md_content += f"\n## Missing Keywords ({len(result['missing'])})\n"
        if result['missing']:
            md_content += "These keywords were found in the job description but are missing or not matching in the resume. Consider incorporating them where relevant:\n\n"
            md_content += "| Keyword / Skill | Recommendation |\n| :--- | :--- |\n"
            for kw in sorted(result['missing']):
                md_content += f"| {kw} | Add to skills or experience bullet points ⚠️ |\n"
        else:
            md_content += "🎉 Excellent! No required keywords from the JD are missing in your resume.\n"
            
        with open(args.markdown_output, 'w', encoding='utf-8') as f:
            f.write(md_content)
        print(f"Saved Markdown report to: {args.markdown_output}")

if __name__ == "__main__":
    main()
